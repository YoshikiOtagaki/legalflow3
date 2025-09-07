"""
Document Generation Service using python-docx
"""

import os
import uuid
from datetime import datetime
from typing import Dict, Any, Optional, List
from pathlib import Path
import tempfile
import shutil

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

from ..models import DocumentGenerationRequest, DocumentType, DocumentFormat
from ..utils.exceptions import DocumentGenerationError
from ..utils.template_processor import TemplateProcessor
from ..utils.placeholder_replacer import PlaceholderReplacer


class DocumentGenerator:
    """ドキュメント生成サービス"""

    def __init__(self):
        self.template_processor = TemplateProcessor()
        self.placeholder_replacer = PlaceholderReplacer()
        self.output_dir = Path(
            os.getenv("DOCUMENT_OUTPUT_DIR", "/tmp/generated_documents")
        )
        self.template_dir = Path(os.getenv("TEMPLATE_DIR", "templates"))

        # 出力ディレクトリを作成
        self.output_dir.mkdir(parents=True, exist_ok=True)

    async def generate_document(
        self, request: DocumentGenerationRequest
    ) -> Dict[str, Any]:
        """
        ドキュメントを生成

        Args:
            request: ドキュメント生成リクエスト

        Returns:
            生成されたドキュメントの情報
        """
        try:
            # テンプレートファイルを取得
            template_path = await self._get_template_path(request.template_id)
            if not template_path.exists():
                raise DocumentGenerationError(
                    f"Template not found: {request.template_id}"
                )

            # テンプレートを読み込み
            doc = Document(template_path)

            # プレースホルダーを置換
            self.placeholder_replacer.replace_placeholders(doc, request.data)

            # ドキュメントタイプ別の処理
            if request.document_type == DocumentType.CONTRACT:
                await self._process_contract_document(doc, request.data)
            elif request.document_type == DocumentType.AGREEMENT:
                await self._process_agreement_document(doc, request.data)
            elif request.document_type == DocumentType.PETITION:
                await self._process_petition_document(doc, request.data)
            elif request.document_type == DocumentType.MOTION:
                await self._process_motion_document(doc, request.data)
            elif request.document_type == DocumentType.BRIEF:
                await self._process_brief_document(doc, request.data)
            elif request.document_type == DocumentType.MEMORANDUM:
                await self._process_memorandum_document(doc, request.data)
            elif request.document_type == DocumentType.LETTER:
                await self._process_letter_document(doc, request.data)
            elif request.document_type == DocumentType.REPORT:
                await self._process_report_document(doc, request.data)

            # ドキュメントを保存
            output_filename = f"{request.template_id}_{uuid.uuid4().hex[:8]}.docx"
            output_path = self.output_dir / output_filename

            doc.save(str(output_path))

            return {
                "file_path": str(output_path),
                "filename": output_filename,
                "file_size": output_path.stat().st_size,
                "created_at": datetime.utcnow().isoformat(),
                "document_type": request.document_type,
                "template_id": request.template_id,
            }

        except Exception as e:
            raise DocumentGenerationError(f"Failed to generate document: {str(e)}")

    async def _get_template_path(self, template_id: str) -> Path:
        """テンプレートファイルのパスを取得"""
        return self.template_dir / f"{template_id}.docx"

    async def _process_contract_document(
        self, doc: Document, data: Dict[str, Any]
    ) -> None:
        """契約書ドキュメントの処理"""
        # 契約書固有の処理
        self._add_contract_elements(doc, data)

    async def _process_agreement_document(
        self, doc: Document, data: Dict[str, Any]
    ) -> None:
        """合意書ドキュメントの処理"""
        # 合意書固有の処理
        self._add_agreement_elements(doc, data)

    async def _process_petition_document(
        self, doc: Document, data: Dict[str, Any]
    ) -> None:
        """請願書ドキュメントの処理"""
        # 請願書固有の処理
        self._add_petition_elements(doc, data)

    async def _process_motion_document(
        self, doc: Document, data: Dict[str, Any]
    ) -> None:
        """動議書ドキュメントの処理"""
        # 動議書固有の処理
        self._add_motion_elements(doc, data)

    async def _process_brief_document(
        self, doc: Document, data: Dict[str, Any]
    ) -> None:
        """要約書ドキュメントの処理"""
        # 要約書固有の処理
        self._add_brief_elements(doc, data)

    async def _process_memorandum_document(
        self, doc: Document, data: Dict[str, Any]
    ) -> None:
        """覚書ドキュメントの処理"""
        # 覚書固有の処理
        self._add_memorandum_elements(doc, data)

    async def _process_letter_document(
        self, doc: Document, data: Dict[str, Any]
    ) -> None:
        """書簡ドキュメントの処理"""
        # 書簡固有の処理
        self._add_letter_elements(doc, data)

    async def _process_report_document(
        self, doc: Document, data: Dict[str, Any]
    ) -> None:
        """報告書ドキュメントの処理"""
        # 報告書固有の処理
        self._add_report_elements(doc, data)

    def _add_contract_elements(self, doc: Document, data: Dict[str, Any]) -> None:
        """契約書要素を追加"""
        # 契約書の表紙
        if "contract_title" in data:
            title_paragraph = doc.add_heading(data["contract_title"], level=1)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 当事者情報テーブル
        if "parties" in data:
            self._add_parties_table(doc, data["parties"])

        # 契約条件
        if "terms" in data:
            self._add_terms_section(doc, data["terms"])

        # 署名欄
        self._add_signature_section(doc, data)

    def _add_agreement_elements(self, doc: Document, data: Dict[str, Any]) -> None:
        """合意書要素を追加"""
        # 合意書の表紙
        if "agreement_title" in data:
            title_paragraph = doc.add_heading(data["agreement_title"], level=1)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 合意事項
        if "agreement_items" in data:
            self._add_agreement_items(doc, data["agreement_items"])

    def _add_petition_elements(self, doc: Document, data: Dict[str, Any]) -> None:
        """請願書要素を追加"""
        # 請願書の表紙
        if "petition_title" in data:
            title_paragraph = doc.add_heading(data["petition_title"], level=1)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 請願理由
        if "petition_reasons" in data:
            self._add_petition_reasons(doc, data["petition_reasons"])

    def _add_motion_elements(self, doc: Document, data: Dict[str, Any]) -> None:
        """動議書要素を追加"""
        # 動議書の表紙
        if "motion_title" in data:
            title_paragraph = doc.add_heading(data["motion_title"], level=1)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 動議内容
        if "motion_content" in data:
            self._add_motion_content(doc, data["motion_content"])

    def _add_brief_elements(self, doc: Document, data: Dict[str, Any]) -> None:
        """要約書要素を追加"""
        # 要約書の表紙
        if "brief_title" in data:
            title_paragraph = doc.add_heading(data["brief_title"], level=1)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 要約内容
        if "brief_content" in data:
            self._add_brief_content(doc, data["brief_content"])

    def _add_memorandum_elements(self, doc: Document, data: Dict[str, Any]) -> None:
        """覚書要素を追加"""
        # 覚書の表紙
        if "memorandum_title" in data:
            title_paragraph = doc.add_heading(data["memorandum_title"], level=1)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 覚書内容
        if "memorandum_content" in data:
            self._add_memorandum_content(doc, data["memorandum_content"])

    def _add_letter_elements(self, doc: Document, data: Dict[str, Any]) -> None:
        """書簡要素を追加"""
        # 宛先
        if "recipient" in data:
            recipient_paragraph = doc.add_paragraph()
            recipient_paragraph.add_run(f"宛先: {data['recipient']}")

        # 件名
        if "subject" in data:
            subject_paragraph = doc.add_paragraph()
            subject_paragraph.add_run(f"件名: {data['subject']}")

        # 本文
        if "body" in data:
            body_paragraph = doc.add_paragraph()
            body_paragraph.add_run(data["body"])

    def _add_report_elements(self, doc: Document, data: Dict[str, Any]) -> None:
        """報告書要素を追加"""
        # 報告書の表紙
        if "report_title" in data:
            title_paragraph = doc.add_heading(data["report_title"], level=1)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 報告内容
        if "report_content" in data:
            self._add_report_content(doc, data["report_content"])

    def _add_parties_table(self, doc: Document, parties: List[Dict[str, Any]]) -> None:
        """当事者情報テーブルを追加"""
        if not parties:
            return

        # テーブルを作成
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # ヘッダー行
        header_cells = table.rows[0].cells
        header_cells[0].text = "当事者名"
        header_cells[1].text = "住所"
        header_cells[2].text = "連絡先"

        # データ行を追加
        for party in parties:
            row_cells = table.add_row().cells
            row_cells[0].text = party.get("name", "")
            row_cells[1].text = party.get("address", "")
            row_cells[2].text = party.get("contact", "")

    def _add_terms_section(self, doc: Document, terms: List[Dict[str, Any]]) -> None:
        """契約条件セクションを追加"""
        if not terms:
            return

        # セクション見出し
        doc.add_heading("契約条件", level=2)

        # 条件を追加
        for i, term in enumerate(terms, 1):
            term_paragraph = doc.add_paragraph()
            term_paragraph.add_run(f"{i}. {term.get('title', '')}").bold = True
            term_paragraph.add_run(f"\n{term.get('description', '')}")

    def _add_signature_section(self, doc: Document, data: Dict[str, Any]) -> None:
        """署名欄を追加"""
        # 署名セクション
        doc.add_heading("署名", level=2)

        # 署名テーブル
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"

        # ヘッダー
        header_cells = table.rows[0].cells
        header_cells[0].text = "当事者A"
        header_cells[1].text = "当事者B"

        # 署名欄
        signature_cells = table.rows[1].cells
        signature_cells[0].text = "\n\n署名: _______________"
        signature_cells[1].text = "\n\n署名: _______________"

    def _add_agreement_items(self, doc: Document, items: List[Dict[str, Any]]) -> None:
        """合意事項を追加"""
        if not items:
            return

        # セクション見出し
        doc.add_heading("合意事項", level=2)

        # 合意事項を追加
        for i, item in enumerate(items, 1):
            item_paragraph = doc.add_paragraph()
            item_paragraph.add_run(f"{i}. {item.get('title', '')}").bold = True
            item_paragraph.add_run(f"\n{item.get('description', '')}")

    def _add_petition_reasons(self, doc: Document, reasons: List[str]) -> None:
        """請願理由を追加"""
        if not reasons:
            return

        # セクション見出し
        doc.add_heading("請願理由", level=2)

        # 理由を追加
        for i, reason in enumerate(reasons, 1):
            reason_paragraph = doc.add_paragraph()
            reason_paragraph.add_run(f"{i}. {reason}")

    def _add_motion_content(self, doc: Document, content: str) -> None:
        """動議内容を追加"""
        if not content:
            return

        # セクション見出し
        doc.add_heading("動議内容", level=2)

        # 内容を追加
        content_paragraph = doc.add_paragraph()
        content_paragraph.add_run(content)

    def _add_brief_content(self, doc: Document, content: str) -> None:
        """要約内容を追加"""
        if not content:
            return

        # セクション見出し
        doc.add_heading("要約内容", level=2)

        # 内容を追加
        content_paragraph = doc.add_paragraph()
        content_paragraph.add_run(content)

    def _add_memorandum_content(self, doc: Document, content: str) -> None:
        """覚書内容を追加"""
        if not content:
            return

        # セクション見出し
        doc.add_heading("覚書内容", level=2)

        # 内容を追加
        content_paragraph = doc.add_paragraph()
        content_paragraph.add_run(content)

    def _add_report_content(self, doc: Document, content: str) -> None:
        """報告内容を追加"""
        if not content:
            return

        # セクション見出し
        doc.add_heading("報告内容", level=2)

        # 内容を追加
        content_paragraph = doc.add_paragraph()
        content_paragraph.add_run(content)

    def _format_document(self, doc: Document, data: Dict[str, Any]) -> None:
        """ドキュメントのフォーマットを調整"""
        # フォント設定
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = "MS Mincho"
                run.font.size = Pt(12)

        # ページ設定
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    async def generate_from_template(
        self,
        template_id: str,
        data: Dict[str, Any],
        output_format: DocumentFormat = DocumentFormat.DOCX,
    ) -> Dict[str, Any]:
        """
        テンプレートからドキュメントを生成

        Args:
            template_id: テンプレートID
            data: ドキュメントデータ
            output_format: 出力フォーマット

        Returns:
            生成されたドキュメントの情報
        """
        request = DocumentGenerationRequest(
            template_id=template_id,
            document_type=DocumentType.OTHER,
            output_format=output_format,
            data=data,
            user_id="system",
        )

        return await self.generate_document(request)

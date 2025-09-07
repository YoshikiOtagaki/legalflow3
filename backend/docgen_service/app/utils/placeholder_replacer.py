"""
Placeholder Replacement Utilities
"""

import re
from typing import Dict, Any
from docx import Document


class PlaceholderReplacer:
    """プレースホルダー置換ユーティリティ"""

    def replace_placeholders(self, doc: Document, data: Dict[str, Any]) -> None:
        """プレースホルダーを置換"""
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                text = run.text
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in text:
                        run.text = text.replace(placeholder, str(value))

        # テーブル内のプレースホルダーも置換
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            text = run.text
                            for key, value in data.items():
                                placeholder = f"{{{{{key}}}}}"
                                if placeholder in text:
                                    run.text = text.replace(placeholder, str(value))
